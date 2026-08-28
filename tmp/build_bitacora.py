from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


REFERENCE = Path(r"C:\Users\Samuel\Desktop\Hotel Casa Antigua\Prompts\PROMPT formato Bitacora.docx")
OUTPUT = Path(r"C:\Users\Samuel\Desktop\MiHotel\BITÁCORA desarrollo local, reservas y facturación 25-27 agosto 2026.docx")


def set_run_font(run, size=11, bold=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_spacing(paragraph, before=0, after=8, line=1.0):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    paragraph.paragraph_format.keep_together = False
    paragraph.paragraph_format.widow_control = True


def add_plain(doc, text, size=11, bold=False, before=0, after=8, keep_with_next=False):
    paragraph = doc.add_paragraph(style="Normal")
    set_paragraph_spacing(paragraph, before, after)
    paragraph.paragraph_format.keep_with_next = keep_with_next
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return paragraph


def add_day(doc, text, page_break=False):
    paragraph = doc.add_paragraph(style="Heading 1")
    set_paragraph_spacing(paragraph, before=12 if not page_break else 0, after=10)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.page_break_before = page_break
    run = paragraph.add_run(text)
    set_run_font(run, size=13, bold=True)
    return paragraph


def add_point(doc, number, title, paragraphs):
    subtitle = add_plain(doc, f"{number}. {title}", bold=True, before=10, after=3, keep_with_next=True)
    subtitle.paragraph_format.widow_control = True
    for idx, text in enumerate(paragraphs):
        p = add_plain(doc, text, before=0, after=8)
        if idx < len(paragraphs) - 1:
            p.paragraph_format.keep_with_next = True


shutil.copy2(REFERENCE, OUTPUT)
doc = Document(OUTPUT)

# Preserve the original section properties while replacing the instructional body.
body = doc._element.body
for child in list(body):
    if child.tag != qn("w:sectPr"):
        body.remove(child)

section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = Inches(1.18)
section.right_margin = Inches(1.18)
section.top_margin = Inches(0.98)
section.bottom_margin = Inches(0.98)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.0

heading = doc.styles["Heading 1"]
heading.font.name = "Calibri"
heading._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
heading._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
heading.font.size = Pt(13)
heading.font.bold = True
heading.paragraph_format.space_before = Pt(12)
heading.paragraph_format.space_after = Pt(10)
heading.paragraph_format.keep_with_next = True

add_plain(doc, "BITÁCORA DE DESARROLLO", size=13, bold=True, after=7)
add_plain(doc, "Proyecto: Sistema de Gestión Hotelera - Hotel Casa Antigua", size=13, bold=True, after=7)
add_plain(doc, "Sistema: MiHotel", size=13, bold=True, after=7)
add_plain(doc, "Tecnologías: ASP.NET Core MVC - MySQL - Visual Studio 2022", size=13, bold=True, after=7)
add_plain(doc, "Periodo documentado: 25/08/2026 al 27/08/2026", size=13, bold=True, after=14)

add_day(doc, "Jornada 1 - 25/08/2026")

add_point(doc, 1, "Análisis del cambio de sistema web a operación local", [
    "Se analizó el cambio de alcance solicitado para MiHotel, debido a que el sistema dejaría de funcionar como una solución alojada en línea y pasaría a utilizarse únicamente dentro de las instalaciones del hotel. A partir de esta decisión se revisó qué funciones dependían de un enfoque web y cuáles continuaban siendo necesarias para la administración interna de reservaciones, habitaciones, huéspedes, pagos y cuentas pendientes.",
    "El objetivo de este análisis fue evitar que la adaptación local se limitara a cambiar la forma de iniciar el programa. También se definió que los procesos internos debían conservar sus validaciones, permisos y trazabilidad, aunque ya no existiera un portal público ni una conexión permanente a Internet. Esto permitió establecer una base coherente para retirar servicios externos sin afectar el funcionamiento cotidiano del hotel."
])

add_point(doc, 2, "Adaptación de la aplicación para funcionamiento local y sin dependencias externas", [
    "Se ajustó la configuración general para que la aplicación pudiera ejecutarse en el equipo local, eliminando del flujo principal las referencias a publicación en línea, acceso de clientes por Internet y automatizaciones que ya no correspondían al nuevo alcance. Asimismo, se retiró la expectativa de emitir documentos fiscales desde el sistema, ya que la facturación continuaría realizándose manualmente en la Agencia Virtual de la SAT.",
    "También se revisó el proceso futuro de distribución del programa y se explicó cómo podría prepararse una instalación autónoma y fuera de línea. Se estableció que una entrega de este tipo deberá incluir la aplicación publicada, la base de datos y sus migraciones, los recursos visuales locales, un procedimiento de respaldo y una forma controlada de iniciar los servicios necesarios. El empaquetado definitivo quedó planteado como una etapa posterior y no como parte de esta implementación."
])

add_point(doc, 3, "Corrección del acceso administrativo al sistema", [
    "Durante las pruebas posteriores al cambio local se detectó que el formulario de inicio de sesión permanecía estático después de ingresar las credenciales administrativas. Se revisó el flujo de autenticación, la cuenta existente y la forma en que la aplicación validaba los datos, ya que el problema impedía comprobar el resto de los módulos aunque el sistema lograra iniciar correctamente.",
    "Se habilitó temporalmente una credencial más sencilla para recuperar el acceso y permitir que el usuario ingresara al panel. La intervención se limitó a restablecer el acceso de la cuenta autorizada; no se intentó recuperar una contraseña cifrada como texto legible. Una vez confirmado el ingreso, quedó disponible la opción normal de cambiar la contraseña desde el propio sistema."
])

add_point(doc, 4, "Restauración visual del formulario de inicio de sesión", [
    "Después de solucionar la autenticación se corrigió la presentación del inicio de sesión para que mantuviera la apariencia utilizada antes de la conversión local. Se conservaron los textos, la distribución y la identidad visual que ya resultaban familiares para el operador, retirando únicamente las funciones relacionadas con el entorno web que ya no debían aparecer.",
    "Este ajuste fue importante porque el cambio de arquitectura no debía provocar una interfaz innecesariamente distinta. Mantener el aspecto anterior redujo la curva de aprendizaje y permitió que el personal continuara utilizando el sistema con la misma referencia visual, mientras la lógica interna quedaba preparada para el nuevo entorno."
])

add_point(doc, 5, "Definición operativa del cierre de estadías y la facturación manual", [
    "Se estudió el riesgo de exigir una factura para marcar todas las reservaciones como terminadas. Aunque esa condición podía parecer una medida de seguridad contra cobros no registrados, en la operación real habría obligado a coordinar en el mismo momento a recepción, que realiza el cobro y el check-out, y a administración, que emite las facturas manualmente en SAT. Además, no todos los huéspedes solicitan factura, incluso cuando han proporcionado información fiscal.",
    "Como solución se separaron tres hechos que no deben confundirse: la salida física del huésped, el pago de la cuenta y la entrega de la factura. La habitación puede liberarse al finalizar el check-out; si el huésped solicitó factura, la reservación pasa a una bandeja de seguimiento fiscal; y si no la solicitó, queda terminada sin generar un pendiente innecesario. Esta decisión evita bloqueos operativos y conserva un control claro sobre las solicitudes reales."
])

add_point(doc, 6, "Estado actual del sistema", [
    "Al finalizar la primera jornada, MiHotel podía ejecutarse nuevamente en el entorno local y el acceso administrativo había sido recuperado. El formulario de inicio de sesión conservaba su apariencia original, mientras las referencias funcionales al alojamiento web y a la facturación automática habían quedado fuera del alcance vigente.",
    "También quedó definido el modelo operativo que guiaría el trabajo siguiente: recepción podría completar la salida y liberar la habitación sin esperar una factura, y administración atendería posteriormente únicamente los documentos realmente solicitados. La instalación completamente empaquetada para otro equipo todavía no fue generada, pero su procedimiento quedó identificado para una fase futura."
])

add_point(doc, 7, "Conclusión", [
    "La primera jornada permitió convertir un cambio general de infraestructura en reglas concretas para el sistema. Además de recuperar el acceso y conservar la presentación conocida por el personal, se aclaró que trabajar localmente no significa perder controles de seguridad ni trazabilidad. La separación entre estadía, pago y factura estableció una base operativa más realista y preparó el desarrollo de las funciones fiscales y de reservaciones que se implementaron posteriormente."
])

add_day(doc, "Jornada 2 - 26/08/2026", page_break=True)

add_point(doc, 1, "Implementación de reservaciones agrupadas por fechas separadas", [
    "Se desarrolló la posibilidad de crear varias estadías relacionadas para un mismo huésped cuando sus fechas no son consecutivas. En lugar de convertirlas en una sola reservación continua, cada periodo se conserva como una reservación independiente con su propia entrada, salida, total, saldo, estado, check-in y check-out. La relación se identifica mediante un grupo, lo que facilita reconocer que las fechas pertenecen a una misma planificación.",
    "La creación se protegió mediante una transacción para evitar que se guarde solamente una parte del grupo si alguna fecha presenta un error. La implementación actual utiliza el mismo huésped y la misma habitación para las estadías creadas juntas, valida la disponibilidad individual y limita la separación a uno o dos días intermedios. Los días que quedan entre las estadías no se bloquean y pueden utilizarse para otras reservaciones."
])

add_point(doc, 2, "Integración de estados, disponibilidad y operaciones individuales dentro del grupo", [
    "Se ajustaron el listado y el detalle de reservas para mostrar una etiqueta de grupo sin sustituir la información de cada estadía. No se creó una reservación madre con control especial sobre las demás; cada fecha mantiene su identidad y avanza individualmente por los estados Pendiente de ingreso, En curso, En check-out, Finalizada o Cancelada.",
    "Esta separación permite realizar check-in y check-out en cada visita, cancelar únicamente la fecha necesaria y liberar la habitación al terminar cada periodo. También se incorporaron validaciones para el cierre de la última estadía, tomando en cuenta las fechas anteriores y el saldo del grupo. Con ello se evita que una acción sobre una fecha cierre o cancele silenciosamente todas las demás."
])

add_point(doc, 3, "Distribución de pagos y conexión con Cuentas por Cobrar", [
    "Se amplió el manejo financiero para aceptar pagos dirigidos a una estadía específica o al conjunto de reservaciones agrupadas. Cuando el pago corresponde a varias fechas, el sistema distribuye el monto entre los saldos aplicables y registra la relación entre cada movimiento y cada reservación. Este mecanismo permite conocer cuánto de un pago corresponde a cada hospedaje y evita saldos negativos o importes aplicados por encima del dinero recibido.",
    "También se permitió que una estadía intermedia finalizara aunque existiera deuda relacionada con fechas futuras del mismo grupo. Si la dueña autoriza una salida con saldo, administración debe dejar una observación y la deuda continúa visible en Cuentas por Cobrar. De esta forma, finalizar la ocupación de la habitación no elimina la obligación financiera ni impide registrar pagos posteriores."
])

add_point(doc, 4, "Desarrollo del registro manual de facturas adjuntas", [
    "Se implementó un módulo para registrar las facturas que administración emite manualmente en la Agencia Virtual de la SAT. El sistema no genera FEL ni se comunica con servicios externos; conserva el archivo PDF y tres datos mínimos que permiten identificarlo: NIT receptor, serie y número de DTE. Se descartó la extracción automática del PDF en esta versión porque un reconocimiento incorrecto podía obligar al operador a revisar y corregir numerosos caracteres, reduciendo la agilidad que se buscaba obtener.",
    "El PDF quedó como fuente documental principal y se incorporó una vista previa en miniatura, similar a la utilizada para el DPI. Esta presentación permite reconocer de inmediato si el documento fue agregado y ofrece acciones para ampliarlo, abrirlo o descargarlo. El módulo se limitó exclusivamente a facturas, retirando las opciones de nota de crédito y nota de débito que no correspondían al objetivo definido."
])

add_point(doc, 5, "Creación del seguimiento de facturas pendientes y alcance por estadías", [
    "Se agregó la opción explícita para indicar durante el check-out si el huésped necesita factura. Cuando la respuesta es afirmativa, la reservación finalizada aparece en Pendientes de factura; cuando es negativa, se cierra sin generar tareas adicionales. La existencia de NIT o datos empresariales no selecciona automáticamente la respuesta, porque esos datos no demuestran que el huésped haya solicitado el documento.",
    "Para grupos se incorporaron tres alcances: toda la reserva agrupada, una estadía específica o varias estadías seleccionadas. Los controles secundarios se muestran únicamente después de elegir el alcance correspondiente, evitando ruido visual. Además, se permitió adjuntar facturas desde el expediente aunque la reservación no se encuentre en la bandeja pendiente, lo que cubre solicitudes tardías y pagos realizados después de la salida."
])

add_point(doc, 6, "Mejoras de visualización y auditoría fiscal", [
    "Se añadieron indicadores compactos en la tabla de reservaciones finalizadas para diferenciar las estadías con factura adjunta y aquellas que todavía esperan el documento. La navegación quedó ordenada en las pestañas Todas, Pendientes de ingreso, En curso, En check-out, Finalizadas, Pendientes de factura y Canceladas, manteniendo el seguimiento fiscal separado del estado de ocupación.",
    "El registro de auditoría se simplificó para mostrar Fecha, Hora, Usuario y Detalle. Se retiraron de la vista los nombres técnicos de acciones y representaciones internas de cambios que podían resultar confusos para personal no especializado. El detalle se redacta en lenguaje natural para identificar quién adjuntó, reemplazó o gestionó un documento sin perder la trazabilidad administrativa."
])

add_point(doc, 7, "Conclusión", [
    "Durante la segunda jornada se integraron dos necesidades relacionadas: administrar visitas separadas de un mismo huésped y conservar las facturas emitidas fuera del sistema. Las reservaciones agrupadas mantuvieron la independencia operativa de cada fecha, mientras los pagos y documentos fiscales adquirieron un alcance verificable. El resultado redujo duplicidad, respetó la forma real de trabajo del hotel y permitió continuar usando la SAT manualmente sin perder el control documental dentro de MiHotel."
])

add_day(doc, "Jornada 3 - 27/08/2026", page_break=True)

add_point(doc, 1, "Simplificación de la decisión de facturación en los procesos de cobro", [
    "Se revisaron los lugares en los que aparecía la decisión fiscal y se eliminó su solicitud durante la creación de una reservación. La pregunta quedó concentrada en el check-out y en el pago total de una Cuenta por Cobrar, que son los momentos en los que el huésped normalmente confirma si necesita el documento. El texto visible se redujo a “¿El huésped necesita factura?” para que el operador pueda comprenderlo sin explicaciones técnicas innecesarias.",
    "También se retiró el botón “Abrir expediente de la reservación y facturas” del flujo de cobro, debido a que la consulta documental puede realizarse posteriormente. Este ajuste disminuyó elementos que no aportaban a la acción inmediata y evitó que el operador confundiera el registro del pago con la administración del expediente fiscal."
])

add_point(doc, 2, "Reestructuración de Cuentas por Cobrar sobre una fuente única", [
    "Se modificó el listado de Cuentas por Cobrar para que trabaje directamente con las mismas reservaciones que utiliza el módulo general. Cada estadía no cancelada con saldo pendiente aparece como una fila, incluso cuando pertenece a un grupo. Esto evita almacenar una segunda copia de la reservación y elimina el riesgo de que dos módulos presenten totales, estados o fechas diferentes.",
    "La tabla adoptó los campos principales de las reservaciones finalizadas: cliente, empresa, habitación, entrada, salida, horas reales, total, pago pendiente, estado y acciones. Se conservaron la búsqueda, el ordenamiento y la paginación, y se retiraron los modos separados “Ver por estadías” y “Ver por cliente”, porque las columnas ordenables ya permiten organizar la información sin duplicar funciones. Los iconos de factura permanecieron únicamente en Reservas, donde aportan mayor valor visual."
])

add_point(doc, 3, "Mejora de legibilidad del historial financiero", [
    "Se simplificó el historial de movimientos para que una persona sin conocimientos de sistemas pueda interpretar la cuenta. Las columnas visibles quedaron como Fecha, Movimiento, Forma de pago, Monto y Estado. Los conceptos técnicos se sustituyeron por descripciones directas como Reserva creada, Pago registrado, Pago anulado y Reembolso registrado.",
    "La fila que representa la creación de la reservación ya no muestra una forma de pago, porque en ese momento todavía no se ha recibido dinero. Las columnas de tipo, descripción técnica y acciones fueron retiradas de esta tabla. El estado se conservó porque permite distinguir la vigencia de pagos, anulaciones y reembolsos, por lo que sí tiene una utilidad operativa real."
])

add_point(doc, 4, "Control para cancelar solicitudes de factura", [
    "Se incorporó una acción que permite retirar una solicitud cuando el huésped informa que ya no desea factura. La opción está disponible desde el seguimiento correspondiente, muestra una confirmación para evitar activaciones accidentales y valida en el servidor que no exista una factura activa asociada antes de cambiar la decisión.",
    "Al confirmarse, la reservación sale de Pendientes de factura y la acción queda registrada en la auditoría. Este control atiende un caso poco frecuente, pero evita que una solicitud permanezca indefinidamente en la bandeja y permite corregir el seguimiento sin borrar documentos ya existentes ni alterar el estado finalizado de la estadía."
])

add_point(doc, 5, "Implementación del reembolso excepcional al cancelar estadías pagadas", [
    "Se resolvió el caso en el que una estadía pagada debe cancelarse después de que otras fechas del grupo ya fueron finalizadas o canceladas. Primero, el sistema intenta trasladar el dinero aplicado hacia otras estadías del mismo grupo que todavía tengan saldo. Si no existe saldo suficiente para recibirlo, la cancelación normal se detiene y muestra el monto que tendría que devolverse al huésped.",
    "Administración puede utilizar la acción “Registrar reembolso y cancelar” únicamente después de entregar físicamente el dinero. MiHotel no realiza la transferencia; conserva el pago original y crea un movimiento separado de reembolso con importe, forma de pago, usuario, fecha, reservación o grupo y referencia al movimiento anterior. Toda la redistribución, el reembolso y la cancelación se ejecutan dentro de una transacción, evitando resultados parciales si ocurre un error."
])

add_point(doc, 6, "Validación integral y actualización de la documentación", [
    "Se compiló el proyecto después de integrar los cambios, obteniendo cero errores. Se mantuvieron veintiséis advertencias que ya existían y que no fueron introducidas por el nuevo flujo de reembolso. Además, el usuario realizó pruebas progresivas durante la implementación y confirmó que la carga de facturas, las reservaciones agrupadas, los pendientes fiscales y las demás funciones trabajadas se comportaban correctamente en el uso observado.",
    "Finalmente se consolidó la documentación de Reservas y Cuentas por Cobrar para reflejar las reglas vigentes, distinguiendo funciones implementadas, decisiones reemplazadas y validaciones futuras. La documentación incluye la operación local, la facturación manual, la fuente única de datos, los pagos agrupados, las cancelaciones y los reembolsos. Esta actualización facilita retomar el proyecto posteriormente sin depender de interpretaciones aisladas de conversaciones anteriores."
])

add_point(doc, 7, "Estado actual del sistema y conclusión", [
    "Al cierre de la tercera jornada, el sistema funciona localmente y el conjunto de funciones para adjuntar y dar seguimiento a facturas puede considerarse completado según las pruebas realizadas. Las reservas individuales y agrupadas conservan su propia operación, Cuentas por Cobrar consulta los mismos datos sin duplicarlos y los pagos posteriores pueden generar una solicitud fiscal aun cuando la estadía ya haya terminado.",
    "El trabajo realizado fortaleció la seguridad financiera sin imponer pasos que dificulten la atención en recepción. Las habitaciones se liberan según la salida real, las deudas continúan visibles, las facturas se gestionan únicamente cuando son necesarias y una cancelación pagada no puede ocultar el dinero recibido. Con esta base, los módulos involucrados quedan preparados para continuar con nuevas funciones o con el futuro empaquetado autónomo del sistema."
])

# Keep all new paragraphs in Calibri and prevent accidental extra outline levels.
for paragraph in doc.paragraphs:
    if paragraph.style.name != "Heading 1":
        paragraph.style = doc.styles["Normal"]
    for run in paragraph.runs:
        if run.font.name != "Calibri":
            set_run_font(run, size=11, bold=bool(run.bold))

doc.core_properties.title = "Bitácora de desarrollo de MiHotel - 25 al 27 de agosto de 2026"
doc.core_properties.subject = "Adaptación local, reservaciones agrupadas, cuentas por cobrar y facturación manual"
doc.core_properties.author = "Hotel Casa Antigua"
doc.core_properties.keywords = "MiHotel, bitácora, reservas, facturación, cuentas por cobrar"
doc.save(OUTPUT)
print(OUTPUT)
