from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


DOCX = Path(r"C:\Users\Samuel\Desktop\MiHotel\BITÁCORA desarrollo local, reservas y facturación 22-27 agosto 2026.docx")


def font(run, size=11, bold=False):
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Calibri")
    rpr.rFonts.set(qn("w:hAnsi"), "Calibri")
    rpr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold


def paragraph(doc, text, *, size=11, bold=False, before=0, after=8, keep=False, style="Normal"):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.keep_with_next = keep
    p.paragraph_format.widow_control = True
    r = p.add_run(text)
    font(r, size, bold)
    return p


def day(doc, label, first=False):
    p = paragraph(doc, label, size=13, bold=True, before=12 if first else 0,
                  after=10, keep=True, style="Heading 1")
    p.paragraph_format.page_break_before = False


def point(doc, number, title, texts):
    paragraph(doc, f"{number}. {title}", bold=True, before=10, after=3, keep=True)
    for i, text in enumerate(texts):
        p = paragraph(doc, text, after=8)
        p.paragraph_format.keep_with_next = i < len(texts) - 1


content = [
    (
        "Jornada 1 - 22/08/2026",
        [
            ("Análisis del cambio de sistema web a operación local", [
                "Se analizó el cambio de alcance solicitado para MiHotel, debido a que la aplicación dejaría de funcionar como una solución alojada en línea y pasaría a utilizarse únicamente dentro de las instalaciones del hotel. La revisión permitió diferenciar las funciones internas que debían conservarse de aquellos procesos pensados para acceso público, disponibilidad por Internet o interacción directa con clientes externos.",
                "El propósito no fue solamente cambiar la manera de iniciar el programa, sino redefinir su arquitectura de acuerdo con la operación real del establecimiento. Se estableció que las reservaciones, habitaciones, huéspedes, pagos, usuarios y cuentas pendientes continuarían funcionando con sus validaciones y permisos, aunque el sistema ya no dependiera de un servidor remoto ni de una conexión permanente a Internet."
            ]),
            ("Definición de los componentes que debían mantenerse", [
                "Se revisaron los módulos existentes para evitar que el cambio local eliminara funciones necesarias para recepción y administración. Los procesos de registro de huéspedes, creación de reservaciones, control de habitaciones, cobros, check-in, check-out y seguimiento financiero se conservaron como parte central de la aplicación.",
                "Esta decisión permitió separar la tecnología de alojamiento del propósito del sistema. Aunque la aplicación ya no se publique en la web, continúa utilizando una estructura interna basada en ASP.NET Core MVC y MySQL. Lo que se retira es la dependencia de servicios externos, no las reglas de negocio que protegen la información del hotel."
            ]),
            ("Retiro del enfoque de facturación automática", [
                "También se descartó la emisión automática de facturas electrónicas desde MiHotel. La administración continuará creando los documentos manualmente en la Agencia Virtual de la SAT, por lo que el programa no necesita conectarse con FEL ni depender de credenciales o servicios de facturación en línea.",
                "El nuevo alcance asignó a MiHotel una responsabilidad más concreta: registrar si el huésped necesita factura y conservar posteriormente el PDF emitido. Esta separación reduce la complejidad de la instalación local y evita que una interrupción de Internet impida completar los procesos internos del hotel."
            ]),
            ("Planteamiento de seguridad para el nuevo entorno", [
                "Se determinó que trabajar de manera local no debía interpretarse como una reducción de controles. Los usuarios continúan ingresando con una cuenta, las acciones sensibles deben respetar permisos y los movimientos financieros o fiscales deben conservar fecha, hora y responsable.",
                "La seguridad se enfocó en impedir que una reservación pueda cerrarse o cancelarse ocultando dinero, sin trasladar toda la responsabilidad a la existencia de una factura. Esto permitió preparar reglas más adecuadas para los procesos que posteriormente se implementarían en check-out, Cuentas por Cobrar y cancelaciones."
            ]),
            ("Estado de la jornada y conclusión", [
                "Al finalizar la primera jornada quedó definido el alcance local de MiHotel y se identificaron los procesos web que ya no serían necesarios. La aplicación conservaría su estructura interna y sus controles, pero dejaría fuera el portal público, el alojamiento en línea y la emisión automática de facturas.",
                "Este análisis proporcionó una base técnica y operativa para las siguientes modificaciones. Definir primero los límites del sistema evitó eliminar funciones útiles por error y permitió que el desarrollo posterior respondiera a la forma en que realmente trabajan recepción, administración y la dueña del hotel."
            ])
        ]
    ),
    (
        "Jornada 2 - 23/08/2026",
        [
            ("Adaptación de la configuración para ejecución local", [
                "Se ajustó la configuración general de la aplicación para que pudiera ejecutarse en el equipo local sin depender de publicación en un servidor externo. Se retiraron del flujo principal las referencias a funciones públicas o conectadas y se verificó que los recursos visuales necesarios estuvieran disponibles desde el propio proyecto.",
                "El objetivo fue mantener una experiencia similar a la existente, pero con una ejecución controlada dentro del hotel. La adaptación dejó preparada la aplicación para trabajar con la base de datos local y continuar utilizando los módulos internos sin requerir que el sistema estuviera alojado en Internet."
            ]),
            ("Diagnóstico del problema de acceso administrativo", [
                "Durante las primeras pruebas se detectó que el formulario de inicio de sesión permanecía estático después de ingresar las credenciales de administración. Debido a este problema no era posible comprobar los demás módulos, aunque la aplicación hubiera iniciado correctamente.",
                "Se revisó el flujo de autenticación, la cuenta registrada y la validación de credenciales. También se aclaró que una contraseña almacenada de forma segura no debe desencriptarse para recuperar su texto original; el procedimiento correcto consiste en restablecerla de manera controlada para el usuario autorizado."
            ]),
            ("Recuperación controlada de la cuenta de administrador", [
                "Se habilitó temporalmente una credencial más sencilla para recuperar el acceso y permitir que el usuario ingresara al panel. Después de realizar nuevas pruebas, el inicio de sesión respondió correctamente y quedó disponible la opción normal de cambiar la contraseña desde el propio sistema.",
                "Esta corrección permitió continuar la revisión funcional sin exponer información sensible ni alterar otras cuentas. El cambio se trató como una solución temporal de acceso y no como una reducción permanente de las medidas de seguridad."
            ]),
            ("Restauración visual del formulario de inicio de sesión", [
                "Una vez solucionado el acceso se restauró la presentación del login para que mantuviera los textos, la distribución y la identidad visual utilizada antes de la conversión local. Solo se excluyeron las funciones que dependían del entorno web y que ya no correspondían al alcance vigente.",
                "Conservar esta apariencia redujo la curva de aprendizaje del personal. El cambio de arquitectura quedó principalmente en la lógica y en la forma de despliegue, mientras la interfaz siguió siendo reconocible para las personas que ya utilizaban MiHotel."
            ]),
            ("Estado de la jornada y conclusión", [
                "Al finalizar la segunda jornada, MiHotel podía ejecutarse localmente y el acceso administrativo había sido recuperado. El formulario de autenticación volvía a comportarse correctamente y mantenía la presentación anterior, sin mostrar funciones eliminadas.",
                "También se explicó el proceso futuro para preparar una instalación autónoma: publicar la aplicación, incluir base de datos y migraciones, conservar recursos locales, definir respaldos y proporcionar un inicio controlado de servicios. El instalador definitivo no fue creado en esta etapa, pero quedaron identificados sus componentes para una fase posterior."
            ])
        ]
    ),
    (
        "Jornada 3 - 24/08/2026",
        [
            ("Análisis del riesgo de exigir factura para finalizar reservaciones", [
                "Se estudió la recomendación de obligar a que todas las reservaciones tuvieran factura antes de marcarlas como terminadas. Aunque esta condición podía parecer una protección contra cobros no registrados, en la práctica habría hecho depender el cierre de dos personas con responsabilidades diferentes: recepción realiza el cobro y el check-out, mientras administración emite la factura manualmente en SAT.",
                "La restricción también habría creado una cola permanente para huéspedes que nunca solicitaron factura. Por esta razón se concluyó que la seguridad financiera debía comprobarse mediante saldos, pagos, permisos y auditoría, en lugar de utilizar la factura como requisito universal para liberar una habitación."
            ]),
            ("Separación entre salida, pago y entrega de factura", [
                "Se definieron como hechos independientes la salida física del huésped, la liquidación de la cuenta y la emisión del documento fiscal. La habitación debe liberarse cuando termina la estadía, aunque administración todavía no haya ingresado la factura en el sistema.",
                "Si existe una deuda autorizada, esta continúa en Cuentas por Cobrar después de la salida. Si el huésped necesita factura, la solicitud continúa en una bandeja fiscal. Esta separación evita que la habitación permanezca ocupada artificialmente y permite que cada responsable complete su trabajo en el momento adecuado."
            ]),
            ("Definición de la decisión explícita de facturación", [
                "Se estableció que el sistema debía preguntar directamente si el huésped necesita factura, sin deducirlo por la existencia de NIT, empresa o información fiscal. Un huésped puede proporcionar esos datos y aun así no solicitar el documento, por lo que ninguna opción debe seleccionarse automáticamente.",
                "El momento principal para tomar esta decisión quedó ubicado en el check-out. También se contempló el pago posterior de una cuenta pendiente, porque un huésped frecuente puede retirarse con autorización, cancelar la deuda días después y solicitar la factura únicamente cuando realiza ese pago."
            ]),
            ("Selección de la información fiscal mínima", [
                "Se evaluó almacenar numerosos datos de la factura y extraerlos automáticamente del PDF. Sin embargo, se observó que los errores de lectura podían obligar al operador a corregir números, letras y símbolos, convirtiendo una función de ayuda en una carga adicional.",
                "Se decidió conservar únicamente NIT receptor, serie, número de DTE y archivo PDF. El documento adjunto funciona como fuente principal para consultar el resto de la información. La extracción automática quedó fuera de esta versión, ya que el ingreso manual de pocos campos resulta justificable y más confiable para la operación local."
            ]),
            ("Estado de la jornada y conclusión", [
                "Al concluir la tercera jornada quedó diseñado el flujo fiscal completo antes de modificarlo: no todas las estadías requieren factura, la habitación se libera de forma independiente y solo las solicitudes reales deben aparecer como pendientes. También se definieron los datos mínimos y el uso del PDF como respaldo principal.",
                "Estas decisiones resolvieron el conflicto entre control y agilidad. Recepción puede cerrar la operación del hospedaje, administración puede emitir el documento posteriormente y el sistema mantiene evidencia suficiente sin exigir tareas innecesarias a cada huésped."
            ])
        ]
    ),
    (
        "Jornada 4 - 25/08/2026",
        [
            ("Implementación de reservaciones agrupadas por fechas separadas", [
                "Se desarrolló la creación de varias estadías relacionadas para un mismo huésped cuando las fechas no son consecutivas. Cada periodo permanece como una reservación independiente con entrada, salida, total, saldo, estado, check-in y check-out propios, mientras una identificación de grupo permite reconocer su relación.",
                "La creación se protegió mediante una transacción para evitar grupos incompletos. La implementación actual utiliza el mismo huésped y la misma habitación, valida cada intervalo por separado y limita la separación a uno o dos días intermedios. Esos días continúan disponibles para otras reservaciones porque no forman parte de ninguna estadía."
            ]),
            ("Integración de disponibilidad y estados individuales", [
                "Se modificaron los listados y detalles para mostrar una etiqueta de grupo sin crear una reservación madre con privilegios sobre las demás. Cada fecha avanza individualmente por Pendiente de ingreso, En curso, En check-out, Finalizada o Cancelada.",
                "Esta estructura permite realizar check-in y check-out en cada visita, liberar la habitación al terminar el periodo correspondiente y cancelar únicamente la estadía necesaria. Una acción sobre una fecha no debe cerrar, modificar ni cancelar silenciosamente las otras reservaciones relacionadas."
            ]),
            ("Desarrollo de pagos aplicados a estadías o grupos", [
                "Se amplió el manejo financiero para registrar pagos dirigidos a una fecha específica o al conjunto de estadías agrupadas. Cuando un pago cubre varias fechas, el sistema distribuye el importe entre los saldos aplicables y conserva la relación entre el movimiento y cada reservación.",
                "La distribución impide aplicar más dinero del recibido y evita saldos negativos. Además, permite conocer cuánto corresponde a cada hospedaje y mantener la trazabilidad cuando posteriormente se corrige, anula o redistribuye un movimiento."
            ]),
            ("Ajustes de check-out para estadías agrupadas y deuda autorizada", [
                "Se permitió que una estadía intermedia finalizara aunque el grupo todavía tuviera saldo correspondiente a fechas futuras. La última fecha requiere validar el estado de las anteriores y la situación financiera general para impedir cierres inconsistentes.",
                "Cuando la dueña autoriza una salida con deuda, administración debe registrar una observación obligatoria. La habitación se libera, la estadía queda finalizada y el saldo continúa en Cuentas por Cobrar. De esta manera, el estado de ocupación no elimina la obligación pendiente."
            ]),
            ("Estado de la jornada y conclusión", [
                "Al finalizar la cuarta jornada, las reservaciones agrupadas podían administrarse sin perder la independencia de cada fecha. La disponibilidad intermedia permanecía libre, los estados se controlaban por estadía y los pagos podían distribuirse de forma trazable.",
                "La implementación resolvió la necesidad de representar visitas separadas de un mismo huésped sin tratarlas como una ocupación continua. También dejó una base financiera adecuada para relacionar cuentas pendientes y facturas con una, varias o todas las fechas del grupo."
            ])
        ]
    ),
    (
        "Jornada 5 - 26/08/2026",
        [
            ("Desarrollo del módulo de facturas adjuntas", [
                "Se implementó el registro de las facturas emitidas manualmente en la Agencia Virtual de la SAT. MiHotel no genera FEL ni se comunica con servicios externos; almacena el PDF junto con NIT receptor, serie y número de DTE para mantener el documento relacionado con la reservación correspondiente.",
                "El módulo se limitó exclusivamente a facturas. Se retiraron las opciones de nota de crédito y nota de débito porque no formaban parte del objetivo operativo. Esta reducción evitó agregar conceptos que podían confundir al personal y mantuvo la herramienta concentrada en la necesidad real del hotel."
            ]),
            ("Incorporación de vista previa y consulta del PDF", [
                "Se agregó una miniatura visible del PDF, siguiendo el comportamiento utilizado para mostrar el DPI. La vista previa permite reconocer rápidamente cuándo una factura fue adjuntada, sin depender únicamente del nombre del archivo o de un mensaje de confirmación.",
                "Desde esta sección el documento puede ampliarse, abrirse o descargarse. El PDF permanece como respaldo principal de la información fiscal, lo que reduce la cantidad de campos manuales y facilita revisar el documento completo cuando sea necesario."
            ]),
            ("Creación de Pendientes de factura y sus indicadores", [
                "Se añadió la pestaña Pendientes de factura dentro de Gestión de Reservas. Una estadía finalizada aparece en esta bandeja únicamente cuando existe una solicitud expresa y todavía no se ha registrado una factura que cubra el alcance correspondiente.",
                "También se incorporaron indicadores compactos en las reservaciones finalizadas para distinguir documentos adjuntos y solicitudes pendientes sin agregar columnas excesivas. La navegación quedó ordenada como Todas, Pendientes de ingreso, En curso, En check-out, Finalizadas, Pendientes de factura y Canceladas."
            ]),
            ("Definición del alcance fiscal para reservaciones agrupadas", [
                "Al adjuntar una factura a un grupo se implementaron tres alternativas: toda la reserva agrupada, una estadía específica o varias estadías seleccionadas. Los controles de fechas aparecen únicamente después de elegir la opción que los necesita, evitando mostrar todas las selecciones al mismo tiempo.",
                "El documento puede registrarse incluso cuando la reservación ya no aparece como pendiente. Esto cubre solicitudes tardías y pagos realizados después de la salida. La relación de alcance permite asociar un solo PDF con las fechas cubiertas sin duplicar innecesariamente el archivo."
            ]),
            ("Estado de la jornada y conclusión", [
                "Al cierre de la quinta jornada, administración podía adjuntar facturas PDF, identificar sus datos mínimos, seleccionar el alcance y consultar una vista previa. Recepción podía declarar la necesidad durante el check-out sin quedar obligada a emitir el documento en ese mismo momento.",
                "La auditoría fiscal se simplificó para mostrar Fecha, Hora, Usuario y Detalle en lenguaje comprensible. Con ello, el flujo quedó alineado con el trabajo manual en SAT y proporcionó control documental sin introducir una integración externa ni sobrecargar la interfaz."
            ])
        ]
    ),
    (
        "Jornada 6 - 27/08/2026",
        [
            ("Simplificación de la decisión fiscal y del historial financiero", [
                "Se eliminó la decisión de facturación durante la creación de la reserva y se concentró en check-out y en el pago total de una Cuenta por Cobrar. El texto visible quedó como “¿El huésped necesita factura?”, y se retiró el botón para abrir el expediente fiscal desde el flujo de cobro porque la consulta puede realizarse posteriormente.",
                "El historial financiero se simplificó a Fecha, Movimiento, Forma de pago, Monto y Estado. Los textos técnicos se sustituyeron por Reserva creada, Pago registrado, Pago anulado y Reembolso registrado. La creación no muestra forma de pago, ya que todavía no representa dinero recibido."
            ]),
            ("Reestructuración de Cuentas por Cobrar sin duplicar reservaciones", [
                "Se modificó Cuentas por Cobrar para consultar directamente las mismas reservaciones del módulo general. Cada estadía no cancelada con saldo pendiente aparece como una fila, incluso cuando pertenece a un grupo, evitando almacenar una segunda copia de fechas, estados o importes.",
                "La tabla adoptó cliente, empresa, habitación, entrada, salida, horas reales, total, pago pendiente, estado y acciones. Se conservaron búsqueda, ordenamiento y paginación, mientras los modos separados Ver por estadías y Ver por cliente fueron retirados porque las columnas ya permiten organizar la información."
            ]),
            ("Control para retirar solicitudes de factura", [
                "Se incorporó una acción para cancelar la solicitud cuando el huésped cambia de opinión. La opción muestra una confirmación, valida en el servidor que no exista una factura activa y registra el cambio en auditoría antes de retirar la reservación de Pendientes de factura.",
                "Aunque se trata de un caso poco frecuente, este control impide que una solicitud permanezca indefinidamente en la bandeja. También evita borrar documentos existentes o modificar el estado finalizado de la estadía para corregir únicamente una decisión fiscal."
            ]),
            ("Implementación del reembolso excepcional en cancelaciones pagadas", [
                "Se resolvió el caso en que una estadía pagada debe cancelarse después de que otras fechas del grupo ya terminaron o fueron canceladas. El sistema intenta trasladar primero el dinero hacia estadías del mismo grupo con saldo. Si no existe saldo suficiente, la cancelación se detiene y muestra el monto que tendría que devolverse.",
                "Administración puede confirmar Registrar reembolso y cancelar después de entregar físicamente el dinero. MiHotel conserva el pago original y crea un movimiento separado con importe, forma de pago, usuario, fecha y referencia anterior. La redistribución, el reembolso y la cancelación se ejecutan en una sola transacción para evitar resultados parciales."
            ]),
            ("Estado final, validación y conclusión general", [
                "Se compiló el proyecto después de los cambios con cero errores; permanecieron veintiséis advertencias anteriores que no fueron introducidas por el nuevo flujo. El usuario realizó pruebas progresivas durante el desarrollo y confirmó que las reservaciones agrupadas, Cuentas por Cobrar, los pendientes fiscales y la carga de facturas funcionaban correctamente en el uso observado.",
                "También se consolidó la documentación de Reservas y Cuentas por Cobrar con las reglas vigentes. Al terminar la sexta jornada, MiHotel funcionaba localmente, las facturas podían considerarse completadas según las pruebas realizadas y las operaciones financieras conservaban trazabilidad. La instalación autónoma definitiva permanece como una etapa futura independiente del trabajo funcional ya terminado."
            ])
        ]
    )
]


doc = Document(DOCX)
body = doc._element.body
for child in list(body):
    if child.tag != qn("w:sectPr"):
        body.remove(child)

section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = Inches(1.18)
section.right_margin = Inches(1.18)
section.top_margin = Inches(0.98)
section.bottom_margin = Inches(0.98)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.0

heading = doc.styles["Heading 1"]
heading.font.name = "Calibri"
heading._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
heading._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
heading.font.size = Pt(13)
heading.font.bold = True
heading.paragraph_format.space_after = Pt(10)
heading.paragraph_format.keep_with_next = True

paragraph(doc, "BITÁCORA DE DESARROLLO", size=13, bold=True, after=7)
paragraph(doc, "Proyecto: Sistema de Gestión Hotelera - Hotel Casa Antigua", size=13, bold=True, after=7)
paragraph(doc, "Sistema: MiHotel", size=13, bold=True, after=7)
paragraph(doc, "Tecnologías: ASP.NET Core MVC - MySQL - Visual Studio 2022", size=13, bold=True, after=7)
paragraph(doc, "Periodo documentado: 22/08/2026 al 27/08/2026", size=13, bold=True, after=14)

for day_index, (label, points) in enumerate(content):
    day(doc, label, first=day_index == 0)
    for point_index, (title, texts) in enumerate(points, start=1):
        point(doc, point_index, title, texts)

for p in doc.paragraphs:
    if p.style.name != "Heading 1":
        p.style = doc.styles["Normal"]
    for r in p.runs:
        if r.text:
            size = 13 if p.text.startswith(("BITÁCORA", "Proyecto:", "Sistema:", "Tecnologías:", "Periodo documentado:")) or p.style.name == "Heading 1" else 11
            font(r, size=size, bold=bool(r.bold))

doc.core_properties.title = "Bitácora de desarrollo de MiHotel - seis jornadas"
doc.core_properties.subject = "Adaptación local, reservaciones agrupadas, facturación, cuentas por cobrar y reembolsos"
doc.save(DOCX)
print(DOCX)
